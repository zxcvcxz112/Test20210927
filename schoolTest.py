from docx import Document
from docx.shared import Pt,RGBColor #字号，颜色
from docx.enum.text import WD_ALIGN_PARAGRAPH#对齐
from docx.oxml.ns import qn #中文字体
import os
from docx.shared import Cm,Inches
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # 导入段落对齐
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT # 导入单元格垂直对齐
import xlrd
listDir = os.listdir()
work = Document()
if '家访表.docx' in listDir:
    work =  Document(os.getcwd()+'\家访表.docx')
    table = work.tables[0]
    #listDir = os.listdir()
if '信息表格.xls' in listDir:
    work1= xlrd.open_workbook('信息表格.xls')
    sheet = work1.sheet_by_index(0)
    for i in range(1, sheet.nrows):
        # print(sheet.cell_value(i, 3))
        table.rows[0].cells[1].text=sheet.cell_value(i, 1)
        table.rows[0].cells[3].text=sheet.cell_value(i, 2)
        table.rows[0].cells[7].text=str(sheet.cell_value(i, 3))
        table.rows[1].cells[1].text=sheet.cell_value(i, 6)
        table.rows[2].cells[1].text=str(sheet.cell_value(i, 5))
        table.rows[2].cells[7].text=str(sheet.cell_value(i, 4))
        work.save(os.getcwd()+"\\"+sheet.cell_value(i, 1)+"家访表.docx")
